#!/usr/bin/env python3
"""
将翻译文本按顺序合并为一个 Word 文档。
"""

from __future__ import annotations

from pathlib import Path
from typing import List, Tuple
from datetime import datetime

from docx import Document
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


INPUT_DIR = Path("docs/mvp4")
OUTPUT_DIR = Path("mvp5")
EPUB_SOURCE = Path("TE20260314.epub")
OUTPUT_FILE = OUTPUT_DIR / f"{EPUB_SOURCE.stem}.docx"
TITLE_PREFIX = "标题："


def read_mvp4_article(filepath: Path) -> Tuple[str, List[str]]:
    """读取单篇翻译文本，返回标题和正文段落。"""
    lines = filepath.read_text(encoding="utf-8").splitlines()
    if not lines:
        raise ValueError(f"空文件: {filepath}")

    first_line = lines[0].strip()
    if not first_line.startswith(TITLE_PREFIX):
        raise ValueError(f"标题格式错误: {filepath}")

    title = first_line[len(TITLE_PREFIX):].strip()
    body_lines = lines[1:]

    paragraphs: List[str] = []
    current_paragraph: List[str] = []

    for line in body_lines:
        stripped = line.strip()
        if stripped:
            current_paragraph.append(stripped)
            continue

        if current_paragraph:
            paragraphs.append("\n".join(current_paragraph))
            current_paragraph = []

    if current_paragraph:
        paragraphs.append("\n".join(current_paragraph))

    return title, paragraphs


def build_document(input_dir: Path, output_file: Path) -> int:
    """按文件名顺序合并所有翻译文本到一个 Word 文档。"""
    files = sorted(input_dir.glob("*.txt"))
    if not files:
        raise FileNotFoundError(f"输入目录中没有 .txt 文件: {input_dir}")

    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "宋体"
    normal_style.font.size = Pt(12)
    heading_style = document.styles["Heading 1"]
    heading_style.font.name = "宋体"
    heading_style.font.size = Pt(16)

    added_count = 0

    for index, filepath in enumerate(files, start=1):
        title, paragraphs = read_mvp4_article(filepath)
        numbered_title = f"{index}、{title}"

        heading = document.add_paragraph(style="Heading 1")
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        heading.add_run(numbered_title)

        for paragraph_text in paragraphs:
            document.add_paragraph(paragraph_text)

        added_count += 1

        if index < len(files):
            page_break = document.add_paragraph()
            page_break.add_run().add_break(WD_BREAK.PAGE)

    output_file.parent.mkdir(parents=True, exist_ok=True)
    actual_output = save_document_with_fallback(document, output_file)
    print(f"已生成 Word 文档: {actual_output}")
    return added_count


def output_file_for_epub(epub_path: Path, output_dir: Path) -> Path:
    """根据 epub 文件名生成默认 Word 输出路径。"""
    return output_dir / f"{epub_path.stem}.docx"


def save_document_with_fallback(document: Document, output_file: Path) -> Path:
    """优先覆盖默认文件；若文件被占用，则自动另存为新文件。"""
    try:
        document.save(output_file)
        return output_file
    except PermissionError:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        fallback = output_file.with_name(f"{output_file.stem}_{timestamp}{output_file.suffix}")
        document.save(fallback)
        return fallback


def main() -> int:
    article_count = build_document(INPUT_DIR, OUTPUT_FILE)
    print(f"合并文章数: {article_count}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
